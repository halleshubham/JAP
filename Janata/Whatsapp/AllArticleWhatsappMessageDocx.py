import http.client
import json
import datetime
import re
import unicodedata
import docx

def getSymbols(i):
	symbols = '0️⃣1️⃣2️⃣3️⃣4️⃣5️⃣6️⃣7️⃣8️⃣9️⃣🔟'
	x=str(i)
	if(len(x)==1):
		return((symbols[3*(int(i))])+(symbols[3*(int(i))+2]))
	else:
		i=str(i)
		a=getSymbols(i[0])
		b=getSymbols(i[1])
		return(a+b)

document = docx.Document("E:/Janata Files/Summary .docx")

author=[]
excerpt=[]
title=[]

authorNumber=4
titleNumber=2
excerptNumber=6

paraSize = len(document.paragraphs)

'''
while(excerptNumber<paraSize):
    title1 = ((document.paragraphs)[titleNumber]).text
    author1 = ((document.paragraphs)[authorNumber]).text
    excerpt1 = ((document.paragraphs)[excerptNumber]).text

    if (author1 == '') or (title1 == '') or (excerpt1 == ''):
        print ("There is some error",titleNumber,authorNumber,excerptNumber)
        titleNumber +=1
        authorNumber+=1
        excerptNumber+=1
    else:
        title2 = ((document.paragraphs)[titleNumber]).text
        author2 = ((document.paragraphs)[authorNumber]).text
        excerpt2 = ((document.paragraphs)[excerptNumber]).text

        title2 = unicodedata.normalize("NFKD",title2)
        author2 = unicodedata.normalize("NFKD",author2)
        excerpt2 = unicodedata.normalize("NFKD",excerpt2)

        title3 = title2.split('. ',1)
        author.append(author2)
        title.append(title3[1])
        excerpt.append(excerpt2)

        authorNumber+=6
        titleNumber+=6
        excerptNumber+=6
		'''

for para in document.paragraphs:
    if (para.text != ''):
        if (len(author)== len(excerpt)==len(title)): 
            title1=para.text
            title1=title1.split('. ',1)
            title.append(unicodedata.normalize("NFKD",title1[1]))
            continue
        if (len(excerpt)==len(author)) and (len(excerpt)!=len(title)):
            author.append(unicodedata.normalize("NFKD",para.text))
            continue
        if (len(author)==len(title)) and (len(excerpt)!=len(author)):
            excerpt.append(unicodedata.normalize("NFKD",para.text))
            continue

print (author)
print (excerpt)
print (title)
print ("Uday")

conn = http.client.HTTPSConnection("janataweekly.org")

headers = {
    'content-type': "application/json",
    'cache-control': "no-cache"
    }

conn.request("GET", "/wp-json/wp/v2/posts?per_page=36")
res = conn.getresponse()
data = json.loads(res.read())

refDateObj = datetime.datetime.strptime(data[0]["date"],'%Y-%m-%dT%H:%M:%S')
ntrF = '📮 *Janata Weekly*\n'
ntrF += 'India\'s oldest socialist magazine!\n\n'
ntrF += 'Vol.75, No. '+ str(datetime.date.today().isocalendar()[1] - 4) +' | '+refDateObj.strftime('%d %B, %Y')+' Issue\n\n'
ntrF += 'Editor: Dr.G.G. Parikh \nAssociate Editor: Neeraj Jain \nManaging Editor: Guddi\n\n'
strF1 = ''
strF2 = ''
strF3 = ''
j=1
l=len(author)
everyDayArticleNumber = len/7
reminder = len%7
'''
for i in range(29,-1,-1):
	print(getAuthorName(data[i]["author"]))
	if (refDateObj.date() - datetime.datetime.strptime(data[i]["date"],'%Y-%m-%dT%H:%M:%S').date()).days <= 1:
		l=l+1
'''

for i in range(l,l-12,-1):
	if (refDateObj.date() - datetime.datetime.strptime(data[i]["date"],'%Y-%m-%dT%H:%M:%S').date()).days <= 3:
		strF1 += getSymbols(j)+" *"+title[j-1]+"*\n"

		b= author[j-1][-1]
		if (b == ' '):
			author[j-1] = author[j-1][:-1]

		a= excerpt[j-1][-1]
		if (a == ' '):
			excerpt[j-1] = excerpt[j-1][:-1]

		strF1 += "\n✒️ _"+author[j-1]+"_\n\n"
		#strF1 += "📋 _"+excerpt[j-1]+"_\n\n"
		
		if (i != (l-11)):
			strF1 += data[i]["link"] +"\n-----------------------------------------------------------\n\n" # + "\n IMAGE: " + \
		else:
			strF1 += data[i]["link"]
		j=j+1
print (strF1)

for i in range(l-12,l-22,-1):
	if (refDateObj.date() - datetime.datetime.strptime(data[i]["date"],'%Y-%m-%dT%H:%M:%S').date()).days <= 3:

		strF2 += getSymbols(j)+" *"+title[j-1]+"*\n"
		

		b= author[j-1][-1]
		if (b == ' '):
			author[j-1] = author[j-1][:-1]

		a= excerpt[j-1][-1]
		if (a == ' '):
			excerpt[j-1] = excerpt[j-1][:-1]

		strF2 += "\n✒️ _"+author[j-1]+"_\n\n"
		#strF2 += "📋 _"+excerpt[j-1]+"_\n\n"
		if (i != (l-21)):
			strF2 += data[i]["link"] +"\n-----------------------------------------------------------\n\n" # + "\n IMAGE: " + \
		else:
			strF2 += data[i]["link"]
		j=j+1

for i in range(l-22,-1,-1):
	if (refDateObj.date() - datetime.datetime.strptime(data[i]["date"],'%Y-%m-%dT%H:%M:%S').date()).days <= 3:

		strF3 += getSymbols(j)+" *"+title[j-1]+"*\n"

		b= author[j-1][-1]
		if (b == ' '):
			author[j-1] = author[j-1][:-1]

		a= excerpt[j-1][-1]
		if (a == ' '):
			excerpt[j-1] = excerpt[j-1][:-1]

		strF3 += "\n✒️ _"+author[j-1]+"_\n\n"
		#strF3 += "📋 _"+excerpt[j-1]+"_\n\n"
		if (i != 1):
			strF3 += data[i]["link"] +"\n-----------------------------------------------------------\n\n" # + "\n IMAGE: " + \
		else:
			strF3 += data[i]["link"]
		j=j+1

ntrF3 = "\n➖➖➖➖➖➖➖➖➖➖➖\n\n📋 *About Janata Weekly :*\nJanata Weekly is an *independent socialist journal*. It has raised its challenging voice of principled dissent against all conduct and practice that is detrimental to the cherished values of nationalism, democracy, secularism and socialism, while upholding the integrity and the ethical norms of healthy journalism. It has the enviable reputation of being the oldest continuously published socialist journal in India."
ntrF3 += "\n\n📢 Oldest socialist weekly of India, is now also on facebook!\n"
ntrF3 +='👍 https://facebook.com/JanataWeekly \n\n'
ntrF3 += "📋 *Subscribe to Hard Copy*\n\nAnnual: Rs. 260 /-\nThree Years : Rs. 750 /-\n\n📲 Guddi: 07738082170\n➖➖➖➖➖➖➖➖➖➖➖\n\n"
ntrF3 += "📬 *To recieve Janata directly to your mailbox*\n\nFill this form: https://janataweekly.org/subscribe/\n\n📬 *Join for WhatsApp version* \n"
ntrF3 += "\n*🔴 Group 1*: https://chat.whatsapp.com/GFy7sR6uV9bD9gt5dlhE7D\n\n*🔴 Group 2*: https://chat.whatsapp.com/Gvp7JM00VvZKMeh6ld302v"

#res = test_string[ : N] + add_string + test_string[N : ] 
N =  ntrF.index('\nVol.75')

ntrF1= ntrF[ : N] + '*Part 1*\n' + ntrF[N : ] + strF1 + ntrF3
ntrF2= ntrF[ : N] + '*Part 2*\n' + ntrF[N : ] + strF2 + ntrF3
ntrF4= ntrF[ : N] + '*Part 3*\n' + ntrF[N : ] + strF3 + ntrF3

f1 = open("whatsappMessage1.txt","w",encoding="UTF-8")
f1.write(ntrF1)
f1.close()

f2 = open("whatsappMessage2.txt","w",encoding="UTF-8")
f2.write(ntrF2)
f2.close()

f3 = open("whatsappMessage3.txt","w",encoding="UTF-8")
f3.write(ntrF4)
f3.close()

