import http.client
import json
import datetime
import re
import unicodedata
import docx

document = docx.Document("E:/Lok/Janata/Summary_Janata_July 26.docx")

author=[]
excerpt=[]
title=[]

authorNumber=4
titleNumber=2
excerptNumber=6

para =document.paragraphs

paraSize = len(document.paragraphs)

for para in document.paragraphs:
    if (para.text != '') and (para.text != ' '):
        if (len(author)== len(excerpt)==len(title)): 
            title1=para.text
            title1=title1.split('. ',1)
            title.append(unicodedata.normalize("NFKD",title1[1]))
            continue
        if (len(excerpt)==len(author)) and (len(excerpt)!=len(title)):
            author.append(unicodedata.normalize("NFKD",para.text))
            continue
        if (len(author)==len(title)) and (len(excerpt)!=len(author)):
            excerpt.append(unicodedata.normalize("NFKD",para.text))
            continue

        '''
while(excerptNumber<paraSize):
    title1 = ((document.paragraphs)[titleNumber]).text
    author1 = ((document.paragraphs)[authorNumber]).text
    excerpt1 = ((document.paragraphs)[excerptNumber]).text

    if (author1 == '') or (title1 == '') or (excerpt1 == ''):
        print ("There is some error",titleNumber,authorNumber,excerptNumber)
        titleNumber +=1
        authorNumber+=1
        excerptNumber+=1
    else:
        title2 = ((document.paragraphs)[titleNumber]).text
        author2 = ((document.paragraphs)[authorNumber]).text
        excerpt2 = ((document.paragraphs)[excerptNumber]).text

        title2 = unicodedata.normalize("NFKD",title2)
        author2 = unicodedata.normalize("NFKD",author2)
        excerpt2 = unicodedata.normalize("NFKD",excerpt2)

        title3 = title2.split('. ',1)
        author.append(author2)
        title.append(title3[1])
        excerpt.append(excerpt2)

        authorNumber+=6
        titleNumber+=6
        excerptNumber+=6
        '''

print (author)
print (excerpt)
print (title)
print ("Uday")

conn = http.client.HTTPSConnection("janataweekly.org")

headers = {
    'content-type': "application/json",
    'cache-control': "no-cache"
    }

conn.request("GET", "/wp-json/wp/v2/posts?per_page=36")
res = conn.getresponse()
data = json.loads(res.read())

refDateObj = datetime.datetime.strptime(data[0]["date"],'%Y-%m-%dT%H:%M:%S')
l=len(author)

imp = "-----------------------------------------------------------\n 📱 _Join Lokayat's English Whatsapp Channel:_\nhttps://chat.whatsapp.com/DxY7U42lGBeHyLJwUYCtn4\n\n"
imp += "📱 _लोकायतचे मराठी/हिंदी व्हाट्सएप चॅनेलला जॉईन करा:_\nhttps://chat.whatsapp.com/DQtYRIIuydR2fMTjTHVcOT\n\n"
imp += "📱 _Like Janata Weekly Facebook page:_\nhttps://www.faceboook.com/janataweekly\n\n"
j=0
i=l
strF1=""

while j<l:
	if (refDateObj.date() - datetime.datetime.strptime(data[i]["date"],'%Y-%m-%dT%H:%M:%S').date()).days <= 3:

		strF1 += "⭕ *"+title[j]+"*\n\n"
		strF1 += "✒️ "+author[j]+"\n\n_"+excerpt[j]+"_\n\n*Read full article:*\n"
		strF1 += data[i]["link"]+"\n\n"+imp
		j=j+1
	i=i-1
print (strF1)

f3 = open("EnglishWhatsappMessagesUsingDocx.txt","w",encoding="UTF-8")
f3.write(strF1)
f3.close()

