import base64
import io
import traceback
from time import strftime
from JAP_Utilities.JAP_Authentication.creds_parser import get_creds
import os
from JAP_Utilities.jap import upload_images,get_authors_list,add_authors,create_post,delete_images,delete_posts, process_base64_inline_images
from JAP_Utilities.summary_parser import get_summary_data
from JAP_Utilities.configs_parser import get_params
from datetime import date, datetime
import mammoth
from mammoth import images
import os
import docx
from bs4 import BeautifulSoup
from multiprocessing import Pool, cpu_count
from slugify import slugify
from docx.shared import Pt
from PIL import Image
import re
from docx.shared import Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

def safe_alignment(para):
    """Safely get paragraph alignment, handling 'end'/'start' values that python-docx doesn't support."""
    try:
        return para.alignment
    except Exception:
        # Fallback: read the raw XML jc val attribute
        pPr = para._p.pPr
        if pPr is not None:
            jc = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
            if jc is not None:
                val = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                # Map raw OOXML values to python-docx alignment int equivalents
                mapping = {'start': 0, 'left': 0, 'center': 1, 'right': 2, 'end': 2, 'both': 3, 'justify': 3}
                return mapping.get(val, None)
        return None

def get_full_para_text(para):
    """Get full paragraph text including hyperlink text.
    
    python-docx's para.text only returns text from direct child runs,
    missing text inside <w:hyperlink> elements. This function extracts
    all text from the paragraph XML including hyperlinks.
    """
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    texts = []
    for child in para._p:
        tag = child.tag
        if tag == f'{ns}r':  # Regular run
            t = child.find(f'{ns}t')
            if t is not None and t.text:
                texts.append(t.text)
        elif tag == f'{ns}hyperlink':  # Hyperlink containing runs
            for run in child.findall(f'{ns}r'):
                t = run.find(f'{ns}t')
                if t is not None and t.text:
                    texts.append(t.text)
    return ''.join(texts)

def checkIfAnyElementIsToBeRemoved(tag, title, author):
    while tag is not None and tag.text == '' and tag.img is None:
        nextTag = tag.find_next_sibling('p')
        tag.decompose()
        tag = nextTag
    return tag

def getMatchingSoupElementWithParaText(full_para_text, soupElement):
    while (full_para_text.strip() != soupElement.text.strip()):
        if (soupElement.li != None):
            while (full_para_text not in soupElement.text):
                soupElement = soupElement.next_sibling
            return soupElement

        if ('http' in soupElement.text):
            try:
                linkText = soupElement.text.split(' http')[0]
                if (linkText in full_para_text):
                    break
            except:
                print("soup Element doesn't contain http element not in para.text:",full_para_text)

        soupElement = soupElement.next_sibling
        if soupElement is None:
            break
        if(soupElement.tr != None):
            break
    return soupElement

def extract_table_formatting(docx):
    table_styles = []

    for table in docx.tables:
        table_style = []
        for row in table.rows:
            row_style = []
            for cell in row.cells:
                cell_style = {}
                shading_elements = cell._element.xpath(".//w:shd")
                if len(shading_elements) > 0:
                    cell_shading = shading_elements[0]
                    cell_fill = cell_shading.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                    if cell_fill:
                        cell_style['background-color'] = f"#{cell_fill}"
                
                # Extract borders (top, bottom, left, right)
                for side in ['top', 'bottom', 'left', 'right']:
                    border_elements = cell._element.xpath(f".//w:tcBorders/w:{side}")
                    if len(border_elements) > 0:
                        border = border_elements[0]
                        color = border.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color")
                        size = border.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz")
                        if color is not None and size is not None:
                            cell_style[f'border-{side}'] = f"{int(size)/8}px solid #{color}"
                
                # Text alignment
                if cell.vertical_alignment == WD_ALIGN_VERTICAL.CENTER:
                    cell_style['vertical-align'] = 'middle'
                elif cell.vertical_alignment == WD_ALIGN_VERTICAL.BOTTOM:
                    cell_style['vertical-align'] = 'bottom'
                else:
                    cell_style['vertical-align'] = 'top'
                
                for paragraph in cell.paragraphs:
                    p_align = safe_alignment(paragraph)
                    if p_align == WD_ALIGN_PARAGRAPH.CENTER or p_align == 1:
                        cell_style['text-align'] = 'center'
                    elif p_align == WD_ALIGN_PARAGRAPH.RIGHT or p_align == 2:
                        cell_style['text-align'] = 'right'
                    else:
                        cell_style['text-align'] = 'left'
                
                # Capture text formatting
                text_elements = []
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text_format = ''
                        if run.bold:
                            text_format += 'font-weight: bold;'
                        if run.italic:
                            text_format += 'font-style: italic;'
                        if run.underline:
                            text_format += 'text-decoration: underline;'
                        p_align = safe_alignment(paragraph)
                        if p_align:
                            align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
                            text_format += f'text-align: {align_map.get(p_align, "left")};'
                        
                        text_elements.append({
                            'text': run.text,
                            'style': text_format
                        })
                
                cell_style['text_elements'] = text_elements
                
                row_style.append(cell_style)
            table_style.append(row_style)
        table_styles.append(table_style)
        
    return table_styles


def extract_image_formatting(docx_doc):
    image_styles = []
    # Build a map from XML element -> Paragraph object once for O(1) lookups
    para_map = {para._p: para for para in docx_doc.paragraphs}

    for shape in docx_doc.inline_shapes:
        if shape.type == 3:  # Inline picture type
            image_style = {}
            if shape.width:
                image_style['width'] = f"{shape.width.cm:.2f}cm"
            if shape.height:
                image_style['height'] = f"{shape.height.cm:.2f}cm"
            
            # Find the parent <w:p> element, then use the matching Paragraph object
            # so that safe_alignment() resolves style-inherited alignment too.
            alignment = 'left'  # default
            try:
                parent = shape._inline.getparent()
                while parent is not None and not parent.tag.endswith('}p'):
                    parent = parent.getparent()
                if parent is not None:
                    para = para_map.get(parent)
                    if para is not None:
                        p_align = safe_alignment(para)
                        align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
                        alignment = align_map.get(p_align, 'left')
            except Exception:
                pass
            
            image_style['alignment'] = alignment
            image_styles.append(image_style)
  
    return image_styles

def convertDocxToHtml(docxFilePath,summaryOfArticle):
    title = summaryOfArticle["article_title"]
    title = title.lower()
    title = title.replace(' ','')

    author = summaryOfArticle["article_author"]
    author = author.lower()
    author = author.replace(' ','')

    article1 = docx.Document(docxFilePath)
    
    table_styles = extract_table_formatting(article1)
    
    image_styles = extract_image_formatting(article1)
                        
    html = ''
    with open(docxFilePath, "rb") as docx_file:
        #result = mammoth.convert_to_html(docx_file, convert_image=convert_image)
        result = mammoth.convert_to_html(docx_file)
        article_content = result.value
        with open("file.html", "w", encoding="utf-8") as file:
            file.write(article_content)
        article_content = article_content.replace('\xad', '')
        article_content = article_content.replace('\u00ad', '')
        article_content = article_content.replace('\N{SOFT HYPHEN}', '')
        soup = BeautifulSoup(article_content,'html.parser')
    bullet = ''
    count = 0
    countofH2 = 0
    cool = soup.contents[0]
    all_soup_ps = soup.find_all('p')
    for para in article1.paragraphs:
        full_para_text = get_full_para_text(para)
        if full_para_text.strip():
                  
            cool = getMatchingSoupElementWithParaText(full_para_text, cool)            
            if (type(cool) == type(None)):
                cool = all_soup_ps[count - countofH2]
            compareSoupString = str(cool.text)
            compareSoupString = compareSoupString.lower()
            compareSoupString = compareSoupString.replace(' ','')
            if (title == compareSoupString):
                coolUpdate = cool.nextSibling
                cool.decompose()
                cool = coolUpdate
                continue

            if (author == compareSoupString):
                coolUpdate = cool.nextSibling
                cool.decompose()
                cool = coolUpdate
                continue

            cool = checkIfAnyElementIsToBeRemoved(cool,title,author)

            if ((cool.tr != None)):
                continue

            if ((cool.li != None) and (full_para_text in cool.text)):
                continue

            if ('http' in cool.text):
                cool = cool.nextSibling
                count += 1
                continue

            if (cool.img != None):
                count+=1
                if (cool.text == None):
                    cool = cool.nextSibling

            if (cool.text == full_para_text or (cool.text.strip() == full_para_text.strip())):
                count += 1

                if (para.paragraph_format.left_indent != None or (len(full_para_text) - len(cool.text.strip()) > 2)):
                    cool['style'] = "padding-left: 40px;"

                if para.runs and para.runs[0].font.size == Pt(14):
                    cool.name = 'h2'
                    countofH2 +=1
                if para.runs and para.runs[0].font.underline == True:
                    newTagUnderline = soup.new_tag("span",style="text-decoration: underline;")
                    cool.insert(1,newTagUnderline)
                p_align = safe_alignment(para)
                if p_align == 1:
                    cool['class'] = "has-text-align-center"
                elif p_align == 2:
                    cool['class'] = "has-text-align-right"
            else:
                print("Text DO NOT match with each other\n Para.text is ",full_para_text,"\nCool.text is",cool.text,"\n\n")
    
    img_idx = 0
    for img in soup.find_all("img"):
        if img_idx < len(image_styles):
            img_style = image_styles[img_idx]
            alignment = img_style.get('alignment', 'left')
            
            # Build inline style for dimensions
            style_parts = []
            if 'width' in img_style:
                style_parts.append(f"width: {img_style['width']}")
            if 'height' in img_style:
                style_parts.append(f"height: {img_style['height']}")
            
            # Apply WordPress alignment classes and display style
            wp_align_map = {'center': 'aligncenter', 'right': 'alignright', 'left': 'alignleft'}
            wp_class = wp_align_map.get(alignment, 'alignleft')
            if alignment == 'center':
                style_parts += ['display: block', 'margin-left: auto', 'margin-right: auto']
            elif alignment == 'right':
                style_parts += ['display: block', 'margin-left: auto']
            
            img['style'] = '; '.join(style_parts)
            img['class'] = img.get('class', []) + [wp_class]
            
            # Also set text-align on the parent <p> so paragraph-level alignment is correct
            parent_p = img.find_parent('p')
            if parent_p is not None and alignment != 'left':
                parent_p['class'] = parent_p.get('class', []) + [f'has-text-align-{alignment}']
                parent_p['style'] = f'text-align: {alignment};'
            
            img_idx += 1
            
    for table_idx, table in enumerate(soup.find_all("table")):
        table['class'] = table.get('class', []) + ['table', 'table-bordered', 'table-striped']
        if table_idx < len(table_styles):
            table_style = table_styles[table_idx]
            for row_idx, row in enumerate(table.find_all("tr")):
                if row_idx < len(table_style):
                    row_style = table_style[row_idx]
                    for col_idx, cell in enumerate(row.find_all(["td", "th"])):
                        if col_idx < len(row_style):
                            cell_style = row_style[col_idx]
                            style_str = '; '.join([f"{k}: {v}" for k, v in cell_style.items() if k != 'text_elements'])
                            cell['style'] = style_str
                            
                            # Apply text formatting within the cell
                            cell.clear()
                            for text_elem in cell_style.get('text_elements', []):
                                span = soup.new_tag('span')
                                span.string = text_elem['text']
                                span['style'] = text_elem['style']
                                cell.append(span)
    return str(soup)

if __name__ == '__main__':

    print("Total CPUs being used: ", cpu_count()) 

    creds = get_creds()
    start_time = datetime.now();
    if creds:

        # getting the parameters from issue_params.json
        params = get_params()

        # getting Summary data
        summaryfile = params['summaryfile']
        summary_data = get_summary_data(summaryfile)
    
        articles_folder_path = params['articles_folder_path']
   
        artilces_files = {}
        for article in os.listdir(articles_folder_path):
            article_number = article.split('-')[0]
            artilces_files[article_number] = article

        # Adding authors
        authors_list = get_authors_list(summary_data)
        authors_ids = add_authors(authors_list,creds)
        print('\n------------------------------------------------------------\n')

        if authors_ids:
            # Uploading images
            images_folder_path = params['images_folder_path']
            image_dict = upload_images(images_folder_path,creds,len(authors_ids))
            

            # initializing publish date
            publish_date = params['publish_date']
            publish_time_hour = '1'

            posts_created = []
            posts_not_created = []

            if image_dict['status']:
                print_edition_articles = params["print_edition_articles"]
                blog_edition_articles = params["blog_edition_articles"]
                image_ids = image_dict['image_ids']
                total_publish_payload = []
                print(image_ids)
                print('\n------------------------------------------------------------\n')

                for i in range(len(artilces_files)):
                    try:
                        artilce_path = articles_folder_path + artilces_files[str(int(i) + 1)]
                        try:
                            article_html = ''
                            article_html = convertDocxToHtml(artilce_path, summary_data[i])
                            article_content = process_base64_inline_images(article_html, creds)
                        except Exception as e:
                            print(f'There is some issue with getting Html From Beautiful Soup: {artilce_path}')
                            print(f'Error: {e}')
                            traceback.print_exc()
                            continue

                        total_articles = len(image_ids)
                        publish_min = str((total_articles + 1) - int(summary_data[i]['article_number']))   #for publishing the articles in reverse order
                        date_str = publish_date + 'T' + publish_time_hour + ':' + publish_min + ':00'
                        article_date = datetime.strptime(date_str,'%Y-%m-%dT%H:%M:%S')  
                        

                        article_title = summary_data[i]['article_title']
                        article_slug = slugify(article_title, to_lower=True)
                        article_excerpt = summary_data[i]['article_excerpt']
                        #article_slug = slugify(article_title,to_lower=True)+"_"+ strftime(datetime.strptime(publish_date,'%Y-%m-%D'),"%d%m%Y") 
                        article_image_id = image_ids[str(int(i) + 1)]
                        article_author_id = authors_ids[i]

                        if (i + 1) >= print_edition_articles[0] and (i + 1) <= print_edition_articles[1]:
                            categories = "669" #id for published category
                        elif (i + 1) >= blog_edition_articles[0] and (i + 1) <= blog_edition_articles[1]:
                             categories = "521" #id for blog category
                        else: 
                            categories = "521" 

                        data = {
                                'status':'draft',
                                'title' : article_title,
                                'content' : article_content,
                                'slug': article_slug,
                                'excerpt' : article_excerpt,
                                'featured_media':article_image_id,
                                'author': article_author_id,
                                'date': datetime.strftime(article_date,'%Y-%m-%dT%H:%M:%S'),
                                'categories' : categories
                            }    
                        total_publish_payload.append((data, creds))

                    except Exception as e:
                        print(e)
                        print("Could not create the draft for the article " + article_title)
                        delete_images(list(image_dict['image_ids'].values()),creds)
                        posts_not_created.append(article_title)
                        break
                if len(posts_not_created) == 0:           
                    with Pool() as pool:
                        results = pool.map(create_post, total_publish_payload)
                    statuses = []
                    article_ids = []
                    for result in results:
                        if result['status']:
                            posts_created.append(result['article_title'])
                            article_ids.append(result['article_id'])
                        else:
                            posts_not_created.append(result['article_title'])
                        statuses.append(result['status'])
                    
                    if not all(statuses):
                        print("Deleteing the images...")
                        delete_images(list(image_dict['image_ids'].values()),creds)
                        print("Deleting the draft posts...")
                        delete_posts(article_ids,creds)

                    print('\n------------------------------------------------------------\n')
                    print('Posts created for ' + str(len(posts_created)) + ' articles :')
                    for i in range(len(posts_created)):
                        print(str(i + 1) + '. ' + posts_created[i])
                    print('\n------------------------------------------------------------\n')

                    if len(posts_created) == len(artilces_files):
                        print('Posts created for all the articles!')
                    else:
                        print('Posts could not be created for ' + str(len(posts_not_created)) + ' articles :')
                        for i in range(len(posts_not_created)):
                            print(str(i + 1) + '. ' + posts_not_created[i]) 
                    
                    end_time = datetime.now()
                    print("Script finished in:"+ str((end_time - start_time).seconds) +' seconds.')
                else:
                    print(image_dict['message'])
                    delete_images(list(image_dict['image_ids'].values()),creds)
            else:
                print("Skipping post creation")
        else:
            print("Unable to create all the authors! Script stopped.")
